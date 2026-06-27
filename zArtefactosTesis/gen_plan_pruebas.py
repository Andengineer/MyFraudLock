"""
Genera el Plan de Pruebas de MyFraudLock en formato Word (.docx)
SIGUIENDO EL ESTÁNDAR IEEE 829-2008.

⚠ Fuente única de verdad:
    - zArtefactosTesis/Entregable_01_Historias_De_Usuario.xlsx
    - zArtefactosTesis/Entregable_02_Casos_De_Prueba.xlsx

Estructura del Excel de Historias:
    Hoja "EPICAS":           col A=ID Epica, col B=Objetivo, col C=Rango HU
    Hoja "Historias de Usuario": col A=HU, col B=Rol, col C=Necesito,
                                  col D=Quiero para, col E=Redacción completa,
                                  col F=Nº Escenario, col G=Criterio, col H=Contexto,
                                  col I=Evento, col J=Resultado.

Estructura del Excel de Casos:
    Hoja "LISTA CP":  col A=CP, col B=Descripción, col C=HU, col D=Nº Escenario,
                      col E=Criterio.
    Hoja "CPNNN":     R1=Título, R2=Autor, R3=Precondiciones, R5=Encabezado tabla,
                      R6..R8=Pasos (#, paso, datos, esperado), R10=HU relacionada,
                      R11=Postcondiciones.
"""
import os
from datetime import datetime
from collections import defaultdict

import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = os.path.dirname(__file__)
XLS_H = os.path.join(BASE, "Entregable_01_Historias_De_Usuario.xlsx")
XLS_C = os.path.join(BASE, "Entregable_02_Casos_De_Prueba.xlsx")
OUT   = os.path.join(BASE, "Plan_de_Pruebas_MyFraudLock.docx")

# ─── Paleta corporativa ─────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x2B, 0x6C, 0xB0)
COLOR_SECONDARY = RGBColor(0x80, 0x5A, 0xD5)
COLOR_DARK      = RGBColor(0x1A, 0x20, 0x2C)
COLOR_MUTED     = RGBColor(0x71, 0x80, 0x96)
HEX_HEADER      = "2B6CB0"
HEX_ALT_ROW     = "EBF8FF"
HEX_LIGHT       = "F7FAFC"
HEX_API_HEADER  = "805AD5"   # morado para CPs de API

# Casos que corresponden a la API REST (HU0009 — EPICA0005)
API_HU = "HU0009"


# ════════════════════════════════════════════════════════════════
#   LECTURA DE LOS EXCEL
# ════════════════════════════════════════════════════════════════
def _str(v) -> str:
    return "" if v is None else str(v).strip()


def cargar_epicas() -> list[dict]:
    wb = openpyxl.load_workbook(XLS_H, data_only=True)
    ws = wb["EPICAS"]
    epicas = []
    for r in range(3, ws.max_row + 1):  # R1 vacía, R2 cabecera
        id_e = _str(ws.cell(r, 1).value)
        if not id_e:
            continue
        epicas.append({
            "id":        id_e,
            "objetivo":  _str(ws.cell(r, 2).value),
            "rango_hu":  _str(ws.cell(r, 3).value),
        })
    return epicas


def cargar_historias() -> dict:
    """Devuelve {HU0001: {rol, necesito, finalidad, redaccion, escenarios=[...]}}"""
    wb = openpyxl.load_workbook(XLS_H, data_only=True)
    ws = wb["Historias de Usuario"]
    hus = {}
    for r in range(3, ws.max_row + 1):
        hu = _str(ws.cell(r, 1).value)
        if not hu:
            continue
        if hu not in hus:
            hus[hu] = {
                "id":         hu,
                "rol":        _str(ws.cell(r, 2).value),
                "necesito":   _str(ws.cell(r, 3).value),
                "finalidad":  _str(ws.cell(r, 4).value),
                "redaccion":  _str(ws.cell(r, 5).value),
                "escenarios": [],
            }
        hus[hu]["escenarios"].append({
            "num":       _str(ws.cell(r, 6).value),
            "criterio":  _str(ws.cell(r, 7).value),
            "contexto":  _str(ws.cell(r, 8).value),
            "evento":    _str(ws.cell(r, 9).value),
            "resultado": _str(ws.cell(r, 10).value),
        })
    return hus


def cargar_casos() -> list[dict]:
    """Lee la LISTA CP y cada hoja CPNNN para obtener pasos completos."""
    wb = openpyxl.load_workbook(XLS_C, data_only=True)
    ws = wb["LISTA CP"]
    casos = []
    for r in range(3, ws.max_row + 1):
        cp = _str(ws.cell(r, 1).value)
        if not cp:
            continue
        caso = {
            "id":         cp,
            "descripcion":_str(ws.cell(r, 2).value),
            "hu":         _str(ws.cell(r, 3).value),
            "num_esc":    _str(ws.cell(r, 4).value),
            "criterio":   _str(ws.cell(r, 5).value),
            "autor":      "",
            "preconds":   "",
            "postconds":  "",
            "titulo":     "",
            "pasos":      [],
        }
        # Detalle del CP
        if cp in wb.sheetnames:
            cws = wb[cp]
            # R1: "Caso de Prueba: CPxxx: TITULO"
            t1 = _str(cws.cell(1, 1).value)
            if ":" in t1:
                caso["titulo"] = t1.split(":", 2)[-1].strip()
            caso["autor"]    = _str(cws.cell(2, 2).value)
            # R3 contiene "Precondiciones: ..."
            pre = _str(cws.cell(3, 1).value)
            if pre.lower().startswith("precondiciones"):
                pre = pre.split(":", 1)[-1].strip()
            caso["preconds"] = pre
            # Pasos: a partir de R6
            for rr in range(6, cws.max_row + 1):
                num   = _str(cws.cell(rr, 1).value)
                paso  = _str(cws.cell(rr, 2).value)
                datos = _str(cws.cell(rr, 3).value)
                esp   = _str(cws.cell(rr, 4).value)
                if num.isdigit() and paso:
                    caso["pasos"].append({
                        "num": num, "paso": paso,
                        "datos": datos, "esperado": esp,
                    })
                elif _str(cws.cell(rr, 1).value).startswith("Postcondiciones"):
                    post = _str(cws.cell(rr, 1).value)
                    caso["postconds"] = post.split(":", 1)[-1].strip()
        casos.append(caso)
    return casos


# ════════════════════════════════════════════════════════════════
#   HELPERS DE ESTILO WORD
# ════════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
        elif level in (0, 1):
            run.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            run.font.color.rgb = COLOR_SECONDARY
        else:
            run.font.color.rgb = COLOR_DARK
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11,
                  color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_table(doc, headers, rows, col_widths=None,
              header_color=HEX_HEADER, alt_color=HEX_ALT_ROW):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    # Cabecera
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Calibri"; r.font.size = Pt(10); r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Cuerpo
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[1 + ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            for line in str(val).split("\n"):
                if p.runs and p.text:
                    p = cell.add_paragraph()
                r = p.add_run(line)
                r.font.name = "Calibri"; r.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                _set_cell_bg(cell, alt_color)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t


# ── Tarjeta de Caso de Prueba ──────────────────────────────────
def add_caso_prueba(doc, caso, hu, header_color=HEX_HEADER):
    # Caja de título
    box = doc.add_table(rows=1, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Cm(4)
    box.columns[1].width = Cm(13)
    c1 = box.rows[0].cells[0]; c2 = box.rows[0].cells[1]
    _set_cell_bg(c1, header_color)
    p = c1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caso["id"])
    r.font.name = "Calibri"; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(c2, HEX_ALT_ROW)
    p = c2.paragraphs[0]
    r = p.add_run(caso["titulo"] or caso["descripcion"])
    r.font.name = "Calibri"; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = COLOR_DARK
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Metadatos
    meta_filas = [
        ("Historia de Usuario", f'{caso["hu"]}  «{hu["rol"]}»  —  {hu["necesito"]}'),
        ("Escenario",           f'#{caso["num_esc"]}  —  {caso["criterio"]}'),
        ("Autor",               caso["autor"] or "—"),
        ("Precondiciones",      caso["preconds"] or "—"),
    ]
    add_kv_table(doc, meta_filas)

    # Tabla de pasos
    if caso["pasos"]:
        pasos_rows = [
            [p["num"], p["paso"], p["datos"], p["esperado"]]
            for p in caso["pasos"]
        ]
        add_table(doc,
            headers=["#", "Paso", "Datos de prueba", "Resultado esperado"],
            rows=pasos_rows,
            col_widths=[1.2, 5.5, 5, 5.5],
        )

    # Postcondiciones
    add_kv_table(doc, [("Postcondiciones", caso["postconds"] or "—")])

    doc.add_paragraph()


def add_kv_table(doc, kv_rows):
    t = doc.add_table(rows=len(kv_rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Cm(4)
    t.columns[1].width = Cm(13.2)
    for i, (label, value) in enumerate(kv_rows):
        c1 = t.rows[i].cells[0]; c2 = t.rows[i].cells[1]
        _set_cell_bg(c1, HEX_LIGHT)
        p = c1.paragraphs[0]
        r = p.add_run(label)
        r.font.name = "Calibri"; r.font.size = Pt(10); r.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c2.paragraphs[0]
        for line in str(value).split("\n"):
            if p.runs and p.text:
                p = c2.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Calibri"; r.font.size = Pt(10)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


# ════════════════════════════════════════════════════════════════
#   CONSTRUCCIÓN DEL DOCUMENTO
# ════════════════════════════════════════════════════════════════
epicas    = cargar_epicas()
historias = cargar_historias()
casos     = cargar_casos()

# Agrupar casos por HU para los apartados
casos_por_hu = defaultdict(list)
for cp in casos:
    casos_por_hu[cp["hu"]].append(cp)
for hu in casos_por_hu:
    casos_por_hu[hu].sort(key=lambda c: int(c["num_esc"]) if c["num_esc"].isdigit() else 0)

doc = Document()

# Márgenes y estilo base
for section in doc.sections:
    section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.0)
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)

# ============================ PORTADA =============================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n\nPLAN DE PRUEBAS DEL PRODUCTO SOFTWARE")
r.font.name = "Calibri"; r.font.size = Pt(28); r.bold = True
r.font.color.rgb = COLOR_PRIMARY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\nMyFraudLock")
r.font.name = "Calibri"; r.font.size = Pt(36); r.bold = True
r.font.color.rgb = COLOR_DARK

add_paragraph(doc,
    "Sistema Antifraude Transaccional basado en Deep Learning Explicable",
    italic=True, size=14, color=COLOR_MUTED,
    align=WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(doc,
    "\nIEEE Std 829-2008  ·  Software Test Plan\n",
    size=11, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_table(doc,
    headers=["Campo", "Valor"],
    rows=[
        ["Documento",        "Plan de Pruebas del Producto Software"],
        ["Producto",         "MyFraudLock — Sistema Antifraude Transaccional"],
        ["Versión",          "1.0"],
        ["Fecha de emisión", datetime.now().strftime("%d/%m/%Y")],
        ["Estado",           "Aprobado para ejecución"],
        ["Estándar",         "IEEE Std 829-2008"],
        ["Trazabilidad",     f"{len(epicas)} épicas · {len(historias)} historias · {len(casos)} casos de prueba"],
        ["Autor de pruebas", casos[0]["autor"] if casos else "—"],
    ],
    col_widths=[5, 11])

doc.add_page_break()

# ============================ 1. INTRO =============================
add_heading(doc, "1. Introducción", 1)
add_paragraph(doc,
    "El presente documento describe el Plan de Pruebas formal del producto "
    "software MyFraudLock, un sistema antifraude transaccional construido "
    "sobre Django REST Framework y un modelo de Deep Learning (DAFD-Net) "
    "con explicabilidad SHAP. El plan se elabora conforme al estándar "
    "IEEE Std 829-2008 y deriva su contenido directamente de los "
    "entregables previos del proyecto: Historias de Usuario (Entregable 01) "
    "y Casos de Prueba (Entregable 02).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "1.1 Propósito", 2)
add_paragraph(doc,
    "Definir el conjunto de actividades, recursos, criterios y casos de "
    "prueba necesarios para verificar que MyFraudLock cumple con los "
    "criterios de aceptación de cada Historia de Usuario aprobada y que el "
    "endpoint REST POST /api/transacciones/ opera correctamente.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "1.2 Alcance", 2)
add_paragraph(doc, "El plan cubre las siguientes épicas y módulos:", bold=True)
for e in epicas:
    add_paragraph(doc, f"• {e['id']} — {e['objetivo']}  ({e['rango_hu']}).")

add_heading(doc, "1.3 Fuera de alcance", 2)
add_paragraph(doc, "• Pruebas de rendimiento, carga y estrés.")
add_paragraph(doc, "• Pruebas de seguridad ofensiva (OWASP Top 10, pen-testing).")
add_paragraph(doc, "• Pruebas de compatibilidad multi-navegador exhaustiva.")
add_paragraph(doc, "• Pruebas de internacionalización (i18n).")

add_heading(doc, "1.4 Referencias", 2)
add_paragraph(doc, "• Entregable 01 — Historias de Usuario (Entregable_01_Historias_De_Usuario.xlsx).")
add_paragraph(doc, "• Entregable 02 — Casos de Prueba (Entregable_02_Casos_De_Prueba.xlsx).")
add_paragraph(doc, "• Documento de Requisitos Funcionales (requisitos_funcionales.md).")
add_paragraph(doc, "• Diagramas de Arquitectura Lógica y Física (zArtefactosTesis/).")
add_paragraph(doc, "• Guía de Testing del API REST (api_testing_guide.md).")
add_paragraph(doc, "• IEEE Std 829-2008 — Standard for Software and System Test Documentation.")

doc.add_page_break()

# ============================ 2. ÉPICAS ===============================
add_heading(doc, "2. Épicas del Producto", 1)
add_paragraph(doc,
    "El producto se organiza en las siguientes épicas funcionales, cada una "
    "agrupando un conjunto cohesivo de Historias de Usuario:")
add_table(doc,
    headers=["Épica", "Objetivo", "Historias"],
    rows=[[e["id"], e["objetivo"], e["rango_hu"]] for e in epicas],
    col_widths=[2.5, 11, 3])

doc.add_page_break()

# ============================ 3. ELEMENTOS A PROBAR ====================
add_heading(doc, "3. Elementos a Probar (Test Items)", 1)
add_paragraph(doc,
    "Los elementos a probar son las funcionalidades implementadas en "
    "MyFraudLock, derivadas de las Historias de Usuario aprobadas. La "
    "siguiente tabla las resume agrupadas por épica:")
filas = []
for e in epicas:
    rango = e["rango_hu"]
    hu_ids = []
    if "a" in rango.lower():
        try:
            partes = rango.lower().replace("hu", "").split("a")
            i0 = int(partes[0]); i1 = int(partes[1])
            hu_ids = [f"HU{n:04d}" for n in range(i0, i1 + 1)]
        except Exception:
            hu_ids = []
    else:
        hu_ids = [rango.strip()]
    for hu_id in hu_ids:
        hu = historias.get(hu_id)
        if not hu:
            continue
        filas.append([
            hu_id,
            hu["rol"],
            hu["necesito"],
            str(len(hu["escenarios"])),
        ])
add_table(doc,
    headers=["HU", "Rol", "Funcionalidad (Necesito)", "# Esc."],
    rows=filas,
    col_widths=[1.8, 2.8, 9.5, 1.7])

doc.add_page_break()

# ============================ 4. CARACTERÍSTICAS =======================
add_heading(doc, "4. Características a Probar", 1)
add_paragraph(doc,
    "Cada característica corresponde a una Historia de Usuario completa con "
    "todos sus escenarios derivados. La prioridad se asigna en función del "
    "impacto operativo:")
prio_alta = {"HU0001","HU0002","HU0003","HU0004","HU0007","HU0009"}
filas = []
for hu_id, hu in historias.items():
    prio = "Alta" if hu_id in prio_alta else "Media"
    tipo = "Integración API" if hu_id == API_HU else "Funcional"
    filas.append([hu_id, hu["necesito"], tipo, prio])
add_table(doc,
    headers=["HU", "Característica", "Tipo", "Prioridad"],
    rows=filas,
    col_widths=[1.8, 10, 3, 2.2])

doc.add_page_break()

# ============================ 5. ESTRATEGIA ============================
add_heading(doc, "5. Estrategia de Pruebas", 1)

add_heading(doc, "5.1 Enfoque", 2)
add_paragraph(doc,
    "Se aplica un enfoque de prueba de caja negra a nivel funcional, "
    "ejecutando los casos desde la interfaz web del sistema y desde un "
    "cliente HTTP (cURL / Postman) contra el endpoint REST. Cada caso de "
    "prueba verifica un único escenario de Historia de Usuario, manteniendo "
    "la trazabilidad HU → Escenario → CP.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "5.2 Criterios de Entrada", 2)
add_paragraph(doc, "• El código fuente está estable en la rama principal.")
add_paragraph(doc, "• Las migraciones de base de datos se han aplicado correctamente.")
add_paragraph(doc, "• El servidor Django arranca sin errores en http://127.0.0.1:8000/.")
add_paragraph(doc, "• Existe al menos un usuario por cada rol del sistema (ADMIN, ANALISTA, EJECUTIVO).")

add_heading(doc, "5.3 Criterios de Salida", 2)
add_paragraph(doc, "• El 100% de los casos de prioridad Alta se han ejecutado.")
add_paragraph(doc, "• El 95% de los casos ejecutados resultan Aprobados.")
add_paragraph(doc, "• Todo defecto Crítico o Bloqueante detectado ha sido corregido y re-probado.")
add_paragraph(doc, "• No existen defectos de severidad Alta pendientes sin plan de corrección.")

add_heading(doc, "5.4 Criterios de Aceptación", 2)
add_paragraph(doc,
    "Un caso de prueba se considera Aprobado cuando el resultado obtenido "
    "coincide con el resultado esperado descrito en el caso. En caso "
    "contrario se registra como Fallido con la evidencia correspondiente "
    "(captura, log o respuesta HTTP) y se reporta como defecto.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "5.5 Suspensión y Reanudación", 2)
add_paragraph(doc,
    "Las pruebas se suspenderán cuando se detecten defectos bloqueantes "
    "que impidan continuar (por ejemplo, imposibilidad de iniciar sesión "
    "o caída del servidor). Se reanudarán una vez resueltos los defectos "
    "y validados los casos previamente bloqueados.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

doc.add_page_break()

# ============================ 6. ENTORNO ==============================
add_heading(doc, "6. Entorno de Pruebas", 1)
add_table(doc,
    headers=["Recurso", "Especificación"],
    rows=[
        ["Sistema operativo",  "Windows 10/11, Linux Ubuntu 22.04 LTS o macOS 13+"],
        ["Lenguaje",            "Python 3.10 — 3.12"],
        ["Framework web",       "Django 5.2.6"],
        ["API REST",            "Django REST Framework"],
        ["Base de datos",       "SQLite (local) / PostgreSQL (producción en Clever Cloud)"],
        ["Motor ML",            "TensorFlow / Keras + SHAP"],
        ["Navegador",           "Google Chrome 120+ / Firefox 121+"],
        ["Cliente HTTP",        "cURL 8.x o Postman 10.x"],
        ["Servidor de pruebas", "Django runserver en 127.0.0.1:8000"],
        ["Datos de prueba",     "seed_data.py (transacciones simuladas)"],
    ],
    col_widths=[5, 11])

doc.add_page_break()

# ============================ 7. ROLES ================================
add_heading(doc, "7. Roles y Responsabilidades", 1)
add_table(doc,
    headers=["Rol", "Responsabilidad principal"],
    rows=[
        ["Líder de Pruebas (QA Lead)",
         "Aprobar el plan, supervisar la ejecución, consolidar el reporte final."],
        ["Analista de Pruebas (Tester)",
         "Ejecutar los casos de prueba, registrar evidencias, reportar defectos."],
        ["Desarrollador",
         "Atender los defectos reportados, generar parches y validar la corrección."],
        ["Stakeholder / Usuario clave",
         "Validar la pertinencia de los criterios de aceptación de cada Historia."],
    ],
    col_widths=[5, 11])

doc.add_page_break()

# ============================ 8. CALENDARIO ===========================
add_heading(doc, "8. Calendario de Ejecución", 1)
add_table(doc,
    headers=["Fase", "Actividades", "Duración estimada"],
    rows=[
        ["1. Preparación",       "Configurar entorno, poblar BD, validar usuarios.",               "1 día"],
        ["2. Pruebas funcionales",f"Ejecutar los {sum(1 for c in casos if c['hu'] != API_HU)} casos sobre la interfaz web.",   "4 días"],
        ["3. Pruebas de API",    f"Ejecutar los {sum(1 for c in casos if c['hu'] == API_HU)} casos sobre /api/transacciones/.","1 día"],
        ["4. Re-pruebas",        "Re-ejecutar los casos fallidos tras corrección.",                "1 día"],
        ["5. Cierre",            "Consolidar resultados, emitir reporte final.",                    "1 día"],
    ],
    col_widths=[3.5, 9.5, 3])

doc.add_page_break()

# ============================ 9. HISTORIAS ============================
add_heading(doc, "9. Historias de Usuario", 1)
add_paragraph(doc,
    "Catálogo de Historias de Usuario aprobadas que sustentan el alcance del "
    "plan, agrupadas por épica. Cada historia define los escenarios "
    "(criterios de aceptación) que serán validados por los casos de prueba "
    "del capítulo 10.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

for ep in epicas:
    add_heading(doc, f"9.{epicas.index(ep)+1}  {ep['id']}  —  {ep['objetivo']}", 2)
    rango = ep["rango_hu"]
    hu_ids = []
    if "a" in rango.lower():
        try:
            partes = rango.lower().replace("hu", "").split("a")
            i0 = int(partes[0]); i1 = int(partes[1])
            hu_ids = [f"HU{n:04d}" for n in range(i0, i1 + 1)]
        except Exception:
            hu_ids = []
    else:
        hu_ids = [rango.strip()]
    for hu_id in hu_ids:
        hu = historias.get(hu_id)
        if not hu:
            continue
        # Encabezado HU
        add_paragraph(doc, f"{hu['id']}  «{hu['rol']}»", bold=True, size=12,
                      color=COLOR_SECONDARY)
        add_paragraph(doc, hu["redaccion"] or
                      f"Como {hu['rol']} necesito {hu['necesito']} para {hu['finalidad']}.",
                      italic=True, color=COLOR_MUTED,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        # Tabla de escenarios
        filas = [
            [e["num"], e["criterio"], e["contexto"], e["evento"], e["resultado"]]
            for e in hu["escenarios"]
        ]
        add_table(doc,
            headers=["#", "Criterio de Aceptación", "Contexto", "Evento", "Resultado"],
            rows=filas,
            col_widths=[0.8, 3.7, 4, 4, 4.5])
        doc.add_paragraph()

doc.add_page_break()

# ============================ 10. CASOS ===============================
add_heading(doc, "10. Casos de Prueba", 1)
add_paragraph(doc,
    "Catálogo completo de casos de prueba, agrupados por Historia de Usuario. "
    "Cada caso de prueba contiene precondiciones, pasos detallados con sus "
    "datos de entrada y resultado esperado, y la postcondición posterior a "
    "su ejecución.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY)

idx_hu = 0
for hu_id in sorted(casos_por_hu):
    hu = historias.get(hu_id)
    if not hu:
        continue
    idx_hu += 1
    titulo = f"10.{idx_hu}  {hu_id}  —  {hu['necesito']}"
    add_heading(doc, titulo, 2)
    color = HEX_API_HEADER if hu_id == API_HU else HEX_HEADER
    for cp in casos_por_hu[hu_id]:
        add_caso_prueba(doc, cp, hu, header_color=color)

doc.add_page_break()

# ============================ 11. MATRIZ ==============================
add_heading(doc, "11. Matriz de Trazabilidad", 1)
add_paragraph(doc,
    "La matriz vincula cada Historia de Usuario con sus escenarios y los "
    "casos de prueba que los verifican, garantizando cobertura completa.")

filas = []
for hu_id in sorted(historias):
    hu = historias[hu_id]
    cps = casos_por_hu.get(hu_id, [])
    cps_text = ", ".join(c["id"] for c in cps) or "—"
    filas.append([
        hu_id,
        hu["necesito"],
        str(len(hu["escenarios"])),
        str(len(cps)),
        cps_text,
    ])
add_table(doc,
    headers=["HU", "Funcionalidad", "# Esc.", "# CP", "Casos de prueba"],
    rows=filas,
    col_widths=[1.7, 6.5, 1.5, 1.5, 6])

doc.add_page_break()

# ============================ 12. REGISTRO ============================
add_heading(doc, "12. Plantilla de Registro de Ejecución", 1)
add_paragraph(doc,
    "El equipo de pruebas registrará la ejecución de cada caso en la "
    "siguiente plantilla, una fila por ejecución:")
add_table(doc,
    headers=["CP", "Fecha", "Tester", "Resultado", "Defecto asociado", "Observaciones"],
    rows=[[cp["id"], "DD/MM/AAAA", "", "Aprobado / Fallido", "—", "—"]
          for cp in casos[:6]] + [["…","…","…","…","…","…"]],
    col_widths=[1.8, 2.5, 2.8, 3, 3, 3.5])

add_paragraph(doc,
    "\nEn caso de Fallido, adjuntar captura de pantalla y respuesta HTTP "
    "cruda (si aplica) al sistema de seguimiento de defectos.",
    italic=True, color=COLOR_MUTED)

doc.add_page_break()

# ============================ 13. RIESGOS =============================
add_heading(doc, "13. Riesgos y Mitigaciones", 1)
add_table(doc,
    headers=["Riesgo", "Impacto", "Probabilidad", "Mitigación"],
    rows=[
        ["No disponibilidad del servidor de pruebas", "Alto",  "Baja",
         "Mantener una copia local del proyecto y datos seed."],
        ["Modelo ML no carga (artefacto ausente)",     "Alto",  "Media",
         "Verificar que api/ml/ contenga el modelo antes de iniciar."],
        ["Cambios de esquema en migraciones",          "Medio", "Media",
         "Ejecutar 'migrate' antes de cada ciclo de prueba."],
        ["Tiempo de respuesta superior al esperado",   "Bajo",  "Media",
         "Documentar como observación sin bloquear la entrega."],
        ["Datos de prueba insuficientes",              "Medio", "Baja",
         "Poblar la BD con seed_data.py al inicio del ciclo."],
    ],
    col_widths=[5, 2.5, 3, 5.5])

doc.add_page_break()

# ============================ 14. APROBACIÓN ==========================
add_heading(doc, "14. Aprobación del Plan", 1)
add_paragraph(doc,
    "Mediante la firma del presente documento, los responsables abajo "
    "indicados aprueban el contenido del Plan de Pruebas y autorizan "
    "su ejecución.")

doc.add_paragraph()
add_table(doc,
    headers=["Rol", "Nombre", "Firma", "Fecha"],
    rows=[
        ["Líder de Pruebas",         "", "", ""],
        ["Líder de Desarrollo",       "", "", ""],
        ["Responsable del Producto",  "", "", ""],
    ],
    col_widths=[4, 4.5, 4, 3.5])

add_paragraph(doc, "\n— Fin del Plan de Pruebas —",
              italic=True, color=COLOR_MUTED,
              align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f"Documento generado: {OUT}")
print(f"  · Épicas:    {len(epicas)}")
print(f"  · Historias: {len(historias)}")
print(f"  · Casos:     {len(casos)}")
